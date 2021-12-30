#!/usr/bin/env python

"""This program allows the conversion from TXT files in a certain format to a DOCX file in a certain format."""

from typing import List, Tuple, Optional
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import docx.enum.text
from docx.shared import Pt, RGBColor, Mm, Cm
import docx.text.paragraph
import docx.text.run
import re
import argparse
import glob
import os.path
from collections import namedtuple


__author__ = "Anton Höß"
__copyright__ = "Copyright 2021"
__credits__ = list()
__license__ = "BSD"
__version__ = "0.1"
__maintainer__ = "Anton Höß"
__email__ = "anton.hoess42@gmail.com"
__status__ = "Development"


class Txt2Docx:
    """The converter class."""

    def __init__(self, filename: str, output: Optional[str] = None) -> None:
        """Initializes the Converter.

        Parameters
        ----------
        filename : str
            The filename of the TXT file to convert.
        output : str, optional
            The output folder to store the converted DOCX files to.
        """

        self._filename = filename
        self._output = output

        # Document
        self._document = Document()

        # Styles
        self._styles = None
        self._define_styles()

        # Read the file
        self._title = None
        self._title_original = None
        self._ref_no = None
        self._capo = None
        self._authors = None
        self._copyright = None
        self._tab_indent = 11.64
        self._text = list()

        # Create the document
        self._read_file()
        self._build_document()
    # end def

    def save(self, filename: Optional[str] = None, force_overwrite: bool = False) -> None:
        """Saves the build DOCX document to the given filename.
        If no filename is provided, the base-filename of the TXT file gets used.

        Parameters
        ----------
        filename : str, optional
            The filename to save the DOCX document to.
        force_overwrite : bool, default False
            Indicates if the output file(s) shall get overwritten if they already exist.
        """

        # Determine filename
        if not filename:
            filename = os.path.basename(self._filename)

            if filename.endswith(".txt"):
                filename = filename[:-4]

            filename += ".docx"
        # end if

        # Set page settings
        self._set_page_settings()

        # Save the document
        if not os.path.exists(self._output):
            os.mkdir(self._output)

        fn_out = os.path.join(self._output, filename)

        if not os.path.exists(fn_out) or force_overwrite:
            self._document.save(fn_out)
        else:
            raise FileExistsError(f"File {fn_out} already exists and therefore not got overwritten.")
    # end def

    def _read_file(self) -> None:
        """Reads the file and stores the parsed values into the class members."""

        # Read file
        f = open(self._filename, "r", encoding="utf-8")
        lines = f.read().splitlines()

        # Read meta information - the assigned value indicates if this key is mandatory or not
        keys = {"TITLE": True, "TITLE_ORIGINAL": False, "REF_NO": False, "CAPO": False,
                "AUTHORS": True, "COPYRIGHT": True, "TAB_INDENT": False}
        values = dict()
        line_no = 0

        for key, mand in keys.items():
            value = None

            if len(lines) >= 0:
                line = lines[0]

                if line.startswith(f"{key}="):
                    line_no += 1

                    value = line[len(f"{key}="):]

                    # Delete line
                    del lines[0]
                # end if
            # end if

            if value is not None and value.strip() != "":
                values[key] = value
            elif mand:
                raise ValueError(f"Key '{key}' not defined in line {line_no}")
            # end if
        # end for

        # self._title = values["TITLE"].strip().replace("ß", "ẞ").upper()  # unsure, if it's a good idea to use "ẞ"
        self._title = values["TITLE"].strip().upper()
        self._title_original = values["TITLE_ORIGINAL"].strip() \
            if values.get("TITLE_ORIGINAL") is not None else self._title_original
        self._ref_no = values.get("REF_NO").strip() \
            if values.get("REF_NO") is not None else None
        self._capo = values["CAPO"].strip() \
            if values.get("CAPO") is not None else self._title_original
        self._authors = values["AUTHORS"].strip()
        self._copyright = values["COPYRIGHT"].strip()
        self._tab_indent = float(values["TAB_INDENT"]) if values.get("TAB_INDENT") is not None else self._tab_indent

        # Read blocks (which are separated by space)
        start_block = False
        block = None

        for line in lines:
            line = line.strip()
            line = line.replace(" ", " ")

            if line:
                if start_block:
                    block += "\n" + line
                else:
                    start_block = True
                    block = line
                # end if

            else:
                if start_block:
                    start_block = False
                    self._text.append(block)
                    block = None
                # end if
            # end if
        # end for
        self._text.append(block)
    # end def

    def _build_document(self) -> None:
        """Builds the DOCX document from the previously parsed meta values and text blocks."""

        # Title
        title = self._title

        if self._title_original is not None:
            title += " # " + self._title_original
        p = self._add_paragraph(text=title, style="title")

        if self._ref_no is not None:
            p.add_run(text="\t")
            run = p.add_run(text=self._ref_no)
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

        # Authors
        authors = self._authors
        if self._capo is not None:
            authors += "\tCapo " + self._capo
        self._add_paragraph(text=authors, style="authors")
        self._add_paragraph(text="", style="empty_line")

        # Text
        for tb, text_block in enumerate(self._text):
            bold_indices = self._get_bold_indices(text_block)

            # Check if all starting and ending marker indices are in ascending order
            total_indices = [index for pair in bold_indices for index in pair]

            if total_indices != sorted(total_indices[:]):
                raise ValueError(f"The bold marker's indices ({total_indices}) are not in ascending order!")
            # end if

            # Manipulate the string to split it into a paragraph per line
            # -----------------------------------------------------------

            # Get newline indices
            newlines_positions = self._find_all_substrings(text_block, "\n")
            new_text_block = text_block.split("\n")

            # Make all first characters capitalized (should work for all cases, even if there's e.g. a "1. ",
            # since in this case nothing will happen and this lines' text always starts with a capital letter
            for b, block in enumerate(new_text_block):
                new_text_block[b] = block[0].upper() + block[1:]

            # Make all first characters after A: or V: capitalized
            for b, block in enumerate(new_text_block):
                new_text_block[b] = re.sub(r"[A|V]: .", lambda match_object: match_object.group(0).upper(), block)

            # If there are any bold parts in the current text block, determine for each line break
            # if it is within a bold area or not. If so, add a bold ending tag to the end of the line
            # and a bold opening tag to the beginning of the next line.
            if len(bold_indices) > 0:
                for np, newlines_position in enumerate(newlines_positions):
                    within_bold_block = False

                    for start, end in bold_indices:
                        if start < newlines_position < end:
                            within_bold_block = True
                            break
                        # end fi
                    # end for

                    if within_bold_block:
                        new_text_block[np] += "</b>"
                        new_text_block[np + 1] = "<b>" + new_text_block[np + 1]
                    # end if
                # end for
            # end if

            # Write the paragraphs and the runs within every one
            # Each line becomes its own paragraph
            for line in new_text_block:
                # Recalculate the bold indices for each line
                bold_indices = self._get_bold_indices(line)

                # Start with an empty paragraph and add runs to it later on
                p = self._add_paragraph(text="", style="text")

                # Handle runs
                pos = 0
                for start, end in bold_indices:
                    if start > pos:
                        p.add_run(line[pos:start])

                    p.add_run(line[start + 3:end]).bold = True
                    pos = end + 4
                # end for

                if pos < len(line):
                    p.add_run(line[pos:])
                # end if
            # end for

            # Add line between blocks
            if tb < len(self._text) - 1:
                self._add_paragraph(text="", style="empty_line")
            # end if
        # end def

        # Copyright
        self._add_paragraph(text="", style="empty_line")
        self._add_paragraph(text=self._copyright, style="empty_line")
    # end def

    def _define_styles(self) -> None:
        """Defines the styles used for the DOCX document."""

        self._styles = self._document.styles

        # Title
        style = self._styles.add_style("title", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(51, 51, 51)

        # Authors
        style = self._styles.add_style("authors", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(8)

        # Text
        style = self._styles.add_style("text", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Copyright
        style = self._styles.add_style("copyright", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(8)

        # Empty Line
        style = self._styles.add_style("empty_line", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(6)
    # end def

    def _set_page_settings(self):
        """Changes the page settings."""

        sections = self._document.sections

        for section in sections:
            # Page size
            section.page_width = Mm(210)
            section.page_height = Mm(270)

            # Margin
            margin = Cm(2.5)
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin
        # end for
    # end def

    @staticmethod
    def _set_paragraph_format(paragraph: docx.text.paragraph.Paragraph) -> None:
        """Sets some formatting attributes to the given paragraph.

        Parameters
        ----------
        paragraph : docx.text.paragraph.Paragraph
            The paragraph to be formatted.
        """

        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1
    # end def

    def _add_paragraph(self, text: str, style: str) -> docx.text.paragraph.Paragraph:
        """Creates a paragraph using the given style (incl. adding tab stops) and adds it to the document.

        Parameters
        ----------
        text : str
            The text to place in the paragraph.
        style : str
            The style to format the paragraph.
            
        Returns
        -------
        paragraph: docx.text.paragraph.Paragraph
            The created paragraph.
        """

        p = self._document.add_paragraph(text=text, style=style)
        self._set_paragraph_format(p)
        tab_stops = p.paragraph_format.tab_stops
        _tab_stop = tab_stops.add_tab_stop(Cm(self._tab_indent))

        return p
    # end def

    @staticmethod
    def _find_all_substrings(search_str: str, sub_str: str) -> List[int]:
        """Searches all (non-overlapping) substrings in a given string.

        Parameters
        ----------
        search_str : str
            The string to search in.
        sub_str : str
            The substring to search for.
            
        Returns
        -------
        substrings: list of str
            The list of found substrings positions.
        """

        start = 0

        indices = list()

        while True:
            start = search_str.find(sub_str, start)

            if start == -1:
                break
            else:
                indices.append(start)
                start += len(sub_str)
            # end if
        # end while

        return indices
    # end def

    @staticmethod
    def _get_bold_indices(text: str) -> List[Tuple[int, int]]:
        """Searches all (non-overlapping) substrings in a given string.

        Parameters
        ----------
        text : str
            The text to search for bold tags.

        Returns
        -------
        bold_indices: list of (int, int)
            List of tuples with start and end tag for each bold block.
        """

        bold_start_indices = [m.start() for m in re.finditer("<b>", text)]
        bold_end_indices = [m.start() for m in re.finditer("</b>", text)]

        if len(bold_start_indices) != len(bold_end_indices):
            raise ValueError(f"The number ({len(bold_start_indices)}) of bold starting tags "
                             f"(<b>) is not equal to the number "
                             f"({len(bold_end_indices)}) of bold ending tags (</b>)")
        # end if
        bold_indices = list(zip(bold_start_indices, bold_end_indices))

        return bold_indices
    # end def
# end class


def main() -> None:


    """The main function which parses the program arguments and performs the conversion of the specified files."""

    def get_base_path_from_wildcard_path(wc_path: str) -> Tuple[str, int]:
        wc_path = os.path.normpath(wc_path)
        parts = wc_path.split(os.sep)

        # Special treatment for Windows drive letters
        drive = os.path.splitdrive(wc_path)[0]

        if parts[0] == drive:
            parts[0] = drive + os.path.sep

        cur_path = ""
        parts_cnt = 0
        for part in parts:
            tmp_path = os.path.join(cur_path, part)

            if not os.path.exists(tmp_path):
                break

            parts_cnt += 1
            cur_path = os.path.join(cur_path, part)
        # end for

        return cur_path, parts_cnt
    # end def

    def str2bool(v):
        if isinstance(v, bool):
            return v
        # end if

        if v.lower() in ('yes', 'true', 't', 'y', '1'):
            return True

        elif v.lower() in ('no', 'false', 'f', 'n', '0'):
            return False

        else:
            raise argparse.ArgumentTypeError('Boolean value expected.')
        # end if
    # end def

    # Set up parser for command line arguments
    parser = argparse.ArgumentParser(description="Convert TXT files to DOCX files.")
    parser.add_argument("paths", type=str, nargs='+',
                        help=r"Filenames or folder base paths. Accepts wildcards as well as "
                             r"\**\ for recursive search. Example: ./**/*.txt.")
    parser.add_argument("--output", type=str, default=".",
                        help="Output folder")
    parser.add_argument("--force_overwrite", type=str2bool, default=False,
                        help="Overwrite existing output files.")
    parser.add_argument("--suppress_error_output", type=str2bool, default=True,
                        help="Suppress the error output (traceback) and only print the error string "
                             "without any further information.")

    # Parse arguments
    args = parser.parse_args()

    FilePath = namedtuple("FilePath", "base_path rel_path filename")
    file_paths = list()

    for path in args.paths:
        # If argument is a path, add it ...
        if os.path.isfile(path):
            p, f = os.path.basename(path)
            file_paths.append(FilePath(p, "", f))
        else:
            # ... otherwise it is a filter using wildcards, so evaluate it and add all matched files
            base_path, base_path_parts_cnt = get_base_path_from_wildcard_path(path)

            for fn in glob.glob(path, recursive=True):
                if os.path.isfile(fn):
                    p, f = os.path.split(fn)
                    file_paths.append(FilePath(base_path, p[len(base_path) + 1:], f))
                # end if
            # end for
        # end if
    # end for

    if len(file_paths) == 0:
        if not args.suppress_error_output:
            raise ValueError(f"no input paths found in \"{args.paths}\"")
        # end if

    # Remove duplicates
    # files = set(files)  # Not working anymore when not using pure filepaths

    # Process all files
    for base_path, rel_path, file in file_paths:
        print(f"Processing file \"{file}\"...", end="")

        try:
            out_path = os.path.join(args.output, rel_path)
            if not os.path.exists(out_path):
                os.makedirs(out_path)
            doc = Txt2Docx(filename=os.path.join(base_path, rel_path, file), output=out_path)
            doc.save(force_overwrite=args.force_overwrite)
            print(" Finished!")
        except Exception as e:
            print(f"\n=> ERROR occurred: {e}")
            if not args.suppress_error_output:
                raise e
            # end if
        # end try
    # end for
# end def


if __name__ == "__main__":
    main()
# end if
