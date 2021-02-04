# pip install python-docx

from typing import Optional
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Cm, Inches
import os
import re


class Txt2Docx:
    def __init__(self, filename: str) -> None:
        self._filename = filename

        # Document
        self._document = Document()

        # Styles
        self._styles = None
        self._define_styles()

        # Read the file
        self._title = None
        self._authors = None
        self._copyright = None
        self._text = list()

        # Create the document
        self._read_file()
        self._build_document()
    # end def

    def _read_file(self):
        # Read file
        f = open(self._filename, "r", encoding="utf-8")
        lines = f.read().splitlines()

        # Read meta information
        keys = ("TITLE", "AUTHORS", "COPYRIGHT")
        values = list()
        line_no = 0

        for key in keys:
            line_no += 1
            value = None

            if len(lines) >= 0:
                line = lines[0]

                if line.startswith(f"{key}="):
                    value = line[len(f"{key}="):]  # values.append(line[len(f"{key}=")])
                # end if
            # end if

            if value:
                values.append(value)
            else:
                raise ValueError(f"Key '{key}' not defined in line {line_no}")
            # end if

            # Delete line
            del lines[0]
        # end for

        self._title, self._authors, self._copyright = values
        self._title = self._title.upper()

        # Read blocks (which are separated by space)
        start_block = False
        block = None

        for line in lines:
            line = line.strip()

            if line:
                if start_block:
                    block += os.linesep + line
                else:
                    start_block = True
                    block = line
                # end if

            else:
                if start_block:
                    start_block = False
                    self._text.append(block)
                    block = None
                # end if
            # end if
        # end for
        self._text.append(block)
    # end def

    def _build_document(self):
        # Create the document
        #####################

        # Title
        p = self._document.add_paragraph(self._title, style="title")
        self._set_paragraph_format(p)

        # Authors
        p = self._document.add_paragraph(self._authors, style="authors")
        self._set_paragraph_format(p)
        p = self._document.add_paragraph("", style="authors_after")
        self._set_paragraph_format(p)

        # Text
        for tb, text_block in enumerate(self._text):
            bold_start_indices = [m.start() for m in re.finditer('<b>', text_block)]
            bold_end_indices = [m.start() for m in re.finditer('</b>', text_block)]

            if len(bold_start_indices) != len(bold_end_indices):
                raise ValueError(f"The number ({len(bold_start_indices)}) of bold starting tags (<b>) is not equal to the number "
                      f"({len(bold_end_indices)}) of bold ending tags (</b>)")
            # end if
            bold_indices = list(zip(bold_start_indices, bold_end_indices))

            # Check if all starting and ending marker indices are in ascending order
            total_indices = [index for pair in bold_indices for index in pair]

            if total_indices != sorted(total_indices[:]):
                raise ValueError(f"The bold marker's indices ({total_indices}) are not in ascending order")
            # end if

            p = self._document.add_paragraph("")
            self._set_paragraph_format(p)

            pos = 0
            for start, end in bold_indices:
                if start > pos:
                    p.add_run(text_block[pos:start])

                p.add_run(text_block[start + 3:end]).bold = True
                pos = end + 4
            # end for

            if pos < len(text_block):
                p.add_run(text_block[pos:])
            # end if

            # Add line between blocks
            if tb < len(self._text) - 1:
                p = self._document.add_paragraph("")
                self._set_paragraph_format(p)
            # end if
        # end def

        # Copyright
        p = self._document.add_paragraph("", style="copyright")
        self._set_paragraph_format(p)
        p = self._document.add_paragraph(self._copyright, style="copyright")
        self._set_paragraph_format(p)
    # end def

    def save(self, filename: Optional[str] = None) -> None:
        # Determine filename
        if not filename:
            filename = self._filename

            if filename.endswith(".txt"):
                filename = filename[:-4]

            filename += ".docx"
        # end if

        # Set page settings
        self._set_page_settings()

        # Save the document
        self._document.save(filename)
    # end def

    def _set_page_settings(self):
        # Changing page settings
        sections = self._document.sections

        for section in sections:
            # Page size
            section.page_width = Inches(8.5)
            section.page_height = Inches(11.)

            # Margin
            margin = Cm(2.5)
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin
        # end for
    # end def

    @staticmethod
    def _set_paragraph_format(paragraph) -> None:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1
    # end def

    def _define_styles(self):
        # Define styles
        ###############
        self._styles = self._document.styles

        # Title
        style = self._styles.add_style("title", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(51, 51, 51)

        # Authors
        style = self._styles.add_style("authors", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(8)
        style = self._styles.add_style("authors_after", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(6)

        # Text normal
        style = self._styles.add_style("text_normal", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Text bold
        style = self._styles.add_style("text_bold", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.font.bold = True

        # Copyright
        style = self._styles.add_style("copyright", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(8)
    # end def
# end class


def main():
    doc = Txt2Docx(filename=r"Dir, Herr Jesus Christ, Ehre, Macht und Herrlichkeit_AP_201013.txt")
    doc.save()
# end def


if __name__ == "__main__":
    main()
# end if
